# import libraries
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from PIL import Image

# Debug
DEBUG = False
def log_print(s):
    if DEBUG:
        print (s)

# program description
print ("Program for faster report preparation.")
print("")
print ("General description:")
print ("Choose pictures, remember that name of the file will be saved as chapter name")
print ("Program will create raport_pic.docx document in location which you choose.")
print ("Program will create main chapter and subchapter(s) with file name and paste picture in each chapter")
print ("Converted picture(s) to constant width 12cm are centered and pasted to separate chapter.")
print("")

# open files and print path
log_print("1. Choose pictures")
root1 = tk.Tk()
root1.withdraw()
pictures_file_path = filedialog.askopenfilename(filetypes = (("Image files *.jpg;*.jpeg;*.png;*.wmf", "*.jpg;*.jpeg;*.png;*.wmf"),("All files", "*.*")),multiple=True)
log_print("You selected " + str(len(pictures_file_path)) + " file(s) \r\n")

# choose location and create result.docx document
log_print("2. Choose location to save raport_pic.docx document")
root2 = tk.Tk()
root2.withdraw()
raport_file_path = filedialog.askdirectory(parent=root2,initialdir="/",title='Please select a directory')
log_print("Directory with results: " + str(raport_file_path))

# creater chapter, add chapter name and picture to chapter
docx_document = Document()
log_print("Document created")
#docx_document.add_heading('Test Results', level=1)
paragraph = docx_document.add_paragraph("Test Results")
# TODO jak dodać dwa style? Działa tylko jeden (heading lub numbered list)
paragraph.style = docx_document.styles['List Number']
paragraph.add_run()
paragraph.style = docx_document.styles['Heading 1']

# for every picture
for path in pictures_file_path:
    # add name to chapter
    file_name, text_name = os.path.splitext(path)
    file_name = file_name[path.rfind('/')+1:len(file_name)]
    # TODO jak dodać dwa style? Działa tylko jeden (heading lub numbered list)
    paragraph = docx_document.add_paragraph(str(file_name))
    paragraph.style = docx_document.styles['Heading 2']

    # open picture
    # TODO poprawić wczytywanie zdjęcia bo jakość 4 razy niższa (wysokość wmf /2, szerokość wmf /2 - dlaczego?)
    original = Image.open(path)
    # save picture as png because docx library don't know wmf format
    original.save("temp.png",dpi=(300,300),quality=100)
    # add centered picture with width 12cm
    docx_document.add_picture("temp.png",width=Inches(12/2.54))
    # center the image
    last_paragraph = docx_document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # remove the temp.png file because is not needed anymore
    os.remove("temp.png")

docx_document.save(str(raport_file_path) + "/" + "raport_pic.docx")
log_print("Done")